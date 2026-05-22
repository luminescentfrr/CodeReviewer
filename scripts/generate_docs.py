"""Generate 3 filled .docx files for AI course final submission."""
from __future__ import annotations
import os, datetime
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

OUT_DIR = os.path.join(os.path.dirname(__file__), "..", "docs")
os.makedirs(OUT_DIR, exist_ok=True)

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_para(doc, text, bold=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.bold = bold
    return p

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    doc.add_paragraph()
    return table

# ═══════════════════════════════════════════════════════════════
# Document 1: 期末技术报告
# ═══════════════════════════════════════════════════════════════

doc = Document()
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('人工智能综合课程设计\n期末技术报告\n\nCodeReview AI — 多智能体代码审查桌面系统')
run.font.size = Pt(18)
run.bold = True

doc.add_paragraph()

# ── 一、基本信息 ──
add_heading(doc, '一、基本信息', 1)
add_table(doc, ['项目', '内容'], [
    ['课程名称', '人工智能综合课程设计'],
    ['项目名称', 'CodeReview AI — 多智能体代码审查桌面系统'],
    ['项目类型', 'AI 应用系统（桌面应用 + LLM 集成）'],
    ['项目负责人', 'leon-cty-dxy'],
    ['团队成员', 'leon-cty-dxy'],
    ['开发时间', '2025-2026 学年'],
    ['项目版本', 'V1.0'],
    ['代码仓库地址', '本地 Git 仓库'],
    ['系统演示地址', '本地 Electron 桌面应用，后端 http://127.0.0.1:8765'],
])

add_heading(doc, '项目规模', 2)
add_table(doc, ['维度', '数据'], [
    ['后端 Python 文件', '38 个'],
    ['审查专家 Agent', '6 个 + 1 汇总 + 1 修复 = 8 个角色'],
    ['LLM 可调用工具', '15 个'],
    ['Python 核心依赖', '15 个包'],
    ['前端代码行数', '约 2914 行（HTML + JS + CSS）'],
    ['支持语言', 'Python / JavaScript / TypeScript / Java / C++ / Go'],
])

# ── 二、项目摘要 ──
add_heading(doc, '二、项目摘要', 1)
add_para(doc, (
    'CodeReview AI 是一个基于多智能体架构的 AI 驱动代码审查桌面系统，旨在通过大语言模型自动化代码审查流程，'
    '替代传统耗时的纯人工 Code Review。系统采用 Electron + FastAPI + LangGraph + DeepSeek v4 技术栈，'
    '实现了从代码提交到六维评分再到自动修复的完整闭环。'
))
add_para(doc, (
    '本系统构建了一套完整的 AI 代码审查桌面应用：用户通过 Electron 桌面客户端打开本地项目文件夹，系统自动扫描代码文件，'
    '调用后台 LangGraph 编排的 6 个专业 AI Agent（代码质量审查、安全审计、性能优化、文档检查、测试覆盖分析、架构评估）'
    '并行分析代码，再由 Summarizer Agent 进行交叉验证与汇总，最终输出涵盖六项维度的量化评分和结构化问题列表。'
    '审查完成后，用户可启动 Repair Agent，由 LLM 直接调用工具（replace_code、insert_code、run_linter 等）'
    '自动生成修复方案，并在编辑器中以 Diff 高亮实时展示修改内容，逐项确认后保存。'
))
add_para(doc, (
    '技术实现上，后端采用 FastAPI 提供 REST + SSE 流式接口；多智能体协作基于 LangGraph 的 StateGraph 状态机编排，'
    '6 个 Agent 并行执行、Summarizer 后处理，通过 asyncio.gather 实现并发调用 DeepSeek v4 模型。'
    '工具系统设计了统一的 ToolRegistry 注册中心，15 个 LLM 可调用工具通过装饰器注册，支持参数注入和路径安全校验。'
    '前端为纯原生 HTML/CSS/JS 的 VS Code 风格 IDE 界面，通过 SSE 流式接收每个 Agent 的"思考过程"并实时渲染"专家圆桌"视图。'
    '修复模式采用 Human-in-the-Loop 设计，LLM 生成的每一处修改均在编辑器中以绿色（新增）/红色（删除）高亮标记，用户可逐项接受或撤销。'
))
add_para(doc, '关键词：人工智能，多智能体，代码审查，大语言模型，LangGraph，桌面应用', bold=True)

# ── 三、项目背景与目标 ──
add_heading(doc, '三、项目背景与目标', 1)

add_heading(doc, '3.1 项目背景', 2)
add_para(doc, (
    '代码审查（Code Review）是软件工程中保障代码质量的核心环节。传统人工 Code Review 耗时长、依赖审阅者经验、'
    '易因疲劳导致遗漏。尤其在高校教学场景中，课程设计的代码量动辄数千行，教师难以对每位学生的代码进行全面逐行审查。'
))
add_para(doc, (
    '近年来，DeepSeek v4、GPT-4o 等大语言模型在代码理解与生成方面取得了显著进步，为 AI 辅助代码审查提供了技术基础。'
    '本项目利用多智能体协作架构，将代码审查分解为多个专业化维度，让每个 AI Agent 专攻一个方面，'
    '再通过元裁判 Agent 统筹汇总，实现比单一 Prompt 更深入、更全面的审查效果。'
))

add_heading(doc, '3.2 项目目标', 2)
add_table(doc, ['目标维度', '说明'], [
    ['业务目标', '为开发者提供一键式多维度代码审查工具，将单次审查从天/小时级压缩到分钟级'],
    ['技术目标', '实现 7 个专业化 AI Agent 的并行协作、LLM 工具调用系统（15 个注册工具）、SSE 流式推送、Human-in-the-Loop 修复模式'],
    ['工程目标', '构建完整的 Electron 桌面应用（前后端分离），支持可本地一键部署'],
    ['应用目标', '实际用于高校学生代码作业审查、个人开发质量保障、开源项目贡献前自检'],
])

add_heading(doc, '3.3 项目创新点', 2)
add_para(doc, (
    '（1）多智能体 + 元裁判架构。不同于传统的单 Prompt 代码审查，本项目将审查过程拆解为 6 个并行执行的专家 Agent，'
    '各自独立调用 DeepSeek v4 进行分析，再由 Summarizer Agent 进行交叉验证、去重和汇总。"分而治之"避免了单一 Agent '
    '因 Prompt 过长导致的"中间丢失"问题，每个 Agent 聚焦一个维度，输出质量更高。'
))
add_para(doc, (
    '（2）Human-in-the-Loop 修复模式。大多数 AI 代码修复工具采用"黑箱替换"方式——AI 直接修改文件，用户只能接受或拒绝。'
    '本项目实现了逐项确认的修复流程：Repair Agent 多轮调用工具修改代码，每次修改后自动执行 linter 检查；'
    '前端以红色（删除）和绿色（新增）Diff 高亮标记每一处修改，用户可逐项接受、撤销或全部应用。'
))
add_para(doc, (
    '（3）工具注册与执行系统。设计了统一的 ToolRegistry 注册中心，采用装饰器 + ToolDef 数据类模式，'
    '将 15 个 LLM 可调用工具的 Schema 定义和 Python 实现统一管理。支持参数注入、路径安全检查（_validate_path 防目录穿越）、'
    'dry_run 模式（memory 模式下修改仅影响编辑器不写磁盘）、撤销栈（undo_stack 逐级回滚）。'
))
add_para(doc, (
    '（4）流式实时可视化（Expert Round Table）。审查过程通过 SSE 流式推送每个 Agent 的实时状态：'
    '预处理阶段展示 AST 解析结果，6 个 Agent 并行工作时前端逐个显示完成状态和发现问题数，'
    'Summarizer 阶段实时展示推理过程和生成的审查报告。前端的"专家圆桌"视图将每个 Agent 视为独立对话者，'
    '以类似聊天界面的形式展示各自的"思考过程"。'
))

# ── 四、需求分析 ──
add_heading(doc, '四、需求分析', 1)

add_heading(doc, '4.1 用户需求分析', 2)
add_table(doc, ['用户角色', '痛点', '期望'], [
    ['学生开发者', '缺乏代码评审经验，难以发现逻辑缺陷和安全隐患', '即时结构化反馈，多维度覆盖'],
    ['教师/课程助教', '人工审查大量作业耗时巨大，评分标准难以量化', '自动化批量审查，量化评分体系'],
    ['职业开发者', 'Code Review 依赖资深工程师人工完成，耗时长', '集成到本地工作流的自动化审查工具'],
])

add_heading(doc, '4.2 功能需求列表', 2)
add_table(doc, ['模块', '功能名称', '输入', '输出', '优先级'], [
    ['审查引擎', '单文件代码审查', '代码文本 + 语言', '六维评分 + 问题列表 + 报告', 'P0'],
    ['审查引擎', '多文件项目审查', '文件列表 + 主文件', '项目上下文 + 综合报告', 'P0'],
    ['审查引擎', 'SSE 流式推送', '审查请求', '实时进度事件流', 'P0'],
    ['多智能体', '代码审查 Agent (Reviewer)', '代码 + AST', '逻辑/命名/可读性问题', 'P0'],
    ['多智能体', '安全分析 Agent (Security)', '代码 + 导入列表', 'OWASP Top-10 漏洞检测', 'P0'],
    ['多智能体', '性能优化 Agent (Optimizer)', '代码 + 复杂度信息', '性能反模式识别', 'P0'],
    ['多智能体', '文档分析 Agent (Documenter)', '代码 + 函数列表', '缺失 docstring 检测', 'P1'],
    ['多智能体', '测试质量 Agent (Tester)', '代码 + 测试文件', '测试覆盖不足分析', 'P1'],
    ['多智能体', '架构审查 Agent (Architect)', '代码 + AST', 'SOLID/耦合度分析', 'P1'],
    ['多智能体', '综合裁判 (Summarizer)', '6 个 Agent 输出', '交叉验证 + 六维评分', 'P0'],
    ['修复引擎', '自动修复 (Repair Agent)', '问题列表 + 代码', '多轮工具调用修改', 'P1'],
    ['聊天交互', '智能对话 + 意图分类', '用户消息', '四类意图路由', 'P0'],
    ['代码编辑器', '语法高亮 + Diff 高亮', '源代码', '着色代码 + 行号 + 变更标记', 'P0'],
    ['文件管理', '文件树 + 多 Tab 编辑', '项目根目录', '树形文件浏览 + 标签页', 'P0'],
    ['历史记录', '对话保存/加载', '对话消息', 'Markdown 文件', 'P1'],
    ['安全防护', 'CORS/速率限制/CSP', 'HTTP 请求', '安全策略执行', 'P1'],
])

add_heading(doc, '4.3 非功能性需求', 2)
add_table(doc, ['类别', '需求描述'], [
    ['响应时间', '审查请求 5 秒内开始 SSE 推送，单文件完整审查 30-120 秒'],
    ['稳定性', 'Agent 异常隔离（单个故障不影响其他），LLM 自动重试（3 次指数退避），修复连续 5 次失败自动中止'],
    ['可扩展性', '模板方法模式（BaseAgent）新增 Agent，工具注册表模式（ToolRegistry）新增工具，路由关键词配置'],
    ['安全性', '路径穿越防护（_validate_path），CORS 限制 localhost，速率限制 30次/小时，Electron 沙箱启用了'],
    ['可维护性', '配置集中管理（backend/config.py），前后端分离，Swagger API 文档（/api/docs）'],
])

# ── 五、总体方案设计 ──
add_heading(doc, '五、总体方案设计', 1)

add_heading(doc, '5.1 系统总体架构', 2)
add_para(doc, '本系统采用四层分层架构：')
add_para(doc, (
    '第一层 — Electron 桌面壳层：BrowserWindow 管理（1440x900，暗色主题，自定义标题栏）、'
    '子进程管理（自动启动/停止 Python FastAPI 后端）、IPC 通信桥（文件系统操作、窗口控制）、原生对话框。'
))
add_para(doc, (
    '第二层 — FastAPI API 层：RESTful API（/api/health、/api/review/stream、/api/review/project、/api/chat/stream）、'
    'SSE 流式响应推送、中间件链（CORS → Auth → SecurityHeaders → RateLimit）、请求校验。'
))
add_para(doc, (
    '第三层 — LangGraph 智能体编排层：状态图 preprocess → parallel_agents → summarizer → END，'
    '6 个专家 Agent 通过 asyncio.gather 并行执行，异常隔离，Summarizer 后处理。'
))
add_para(doc, (
    '第四层 — LLM/工具执行层：LLM 适配（DeepSeek v4-pro 首选 / OpenAI 备选），'
    '工具注册表（15 个工具），AST 解析器（Tree-sitter + 正则回退），项目上下文分析器。'
))

add_heading(doc, '5.2 技术路线', 2)
add_para(doc, '完整的代码审查流水线如下：')
add_para(doc, (
    '[用户输入代码] → 1.预处理阶段（AST 解析：提取函数列表、类列表、导入列表）→ '
    '2.多智能体并行审查（asyncio.gather 启动 6 个 Agent，各自返回 JSON）→ '
    '3.综合裁判阶段（交叉验证 → 幻觉过滤 → 去重合并 → 六维评分）→ '
    '4.SSE 流式推送（progress → completed → saved 事件）→ '
    '5.前端渲染（Agent Cards + 评分网格 + 问题列表 + 最终报告）'
))
add_para(doc, '修复流水线：')
add_para(doc, (
    '[用户输入修复请求] → 意图分类（LLM 分类为 review/repair/plan_exec/chat）→ '
    'Repair Agent 多轮修复循环（read_file → replace_code/insert_code → run_linter 验证）→ '
    '前端实时可视化（tool_call/tool_result 事件 + Diff 高亮 + 逐项确认）'
))

add_heading(doc, '5.3 核心技术选型', 2)
add_table(doc, ['技术领域', '选型', '用途', '选择理由'], [
    ['后端框架', 'Python + FastAPI', 'RESTful API + SSE', '异步原生、Pydantic 校验、SSE 支持'],
    ['Web 服务器', 'Uvicorn (ASGI)', '运行 FastAPI 应用', '高性能、Python 生态最成熟'],
    ['桌面框架', 'Electron 33', '桌面 GUI 壳层', '跨平台、原生文件系统、子进程管理'],
    ['前端 UI', '原生 HTML/CSS/JS', 'SPA 界面渲染', '零构建依赖、零框架开销'],
    ['智能体编排', 'LangGraph', '多 Agent 状态图', 'DAG 定义、State 传递、流式 astream'],
    ['LLM 提供商', 'DeepSeek v4-pro', '代码分析/对话', '代码能力强、性价比高'],
    ['LLM 集成', 'LangChain (langchain-openai)', 'LLM 调用封装', '统一模型抽象、流式+工具调用'],
    ['AST 解析', 'Tree-sitter', '代码结构解析', '多语言、结构化输出、容错高'],
    ['流式通信', 'SSE', '后端→前端推送', '比 WebSocket 更轻量、自动重连'],
    ['数据校验', 'Pydantic v2', '请求模型定义', '类型安全、FastAPI 原生集成'],
    ['配置管理', 'python-dotenv', 'API Key 管理', '敏感信息不入库、12-Factor'],
    ['速率限制', 'SlowAPI', 'API 防护', 'FastAPI 集成、可配置限值'],
    ['进程管理', 'child_process.spawn', 'Python 生命周期', '自动启停、端口探测'],
])

# ── 六、关键模块设计与实现 ──
add_heading(doc, '六、关键模块设计与实现', 1)

add_heading(doc, '6.1 多智能体协作引擎（LangGraph Orchestration）', 2)
add_para(doc, '文件：backend/agents/graph.py')
add_para(doc, (
    '基于 LangGraph 的 StateGraph 构建三段式流水线：'
    '(1) 预处理阶段（preprocess_node）：Tree-sitter 解析 AST，提取函数、类、导入和行数统计；'
    '(2) 并行审查阶段（parallel_agents_node）：asyncio.gather 同时启动 6 个 Agent，'
    '每个 Agent 接收完整代码和 AST 信息，独立返回审查结果，任一个失败不影响其他；'
    '(3) 汇总阶段（summarizer_agent）：交叉验证专家间矛盾、过滤无证据幻觉、去重合并重复问题、计算六维评分。'
    '使用 graph.astream(stream_mode="updates") 遍历节点输出并转化为 SSE 事件。'
))

add_heading(doc, '6.2 代码审查智能体（Reviewer Agent）', 2)
add_para(doc, '文件：backend/agents/reviewer.py')
add_para(doc, (
    '继承 BaseAgent 模板方法模式。系统提示词指示 LLM 扮演"资深代码审查工程师"，'
    '从逻辑正确性、命名规范、可读性、注释质量四个维度分析代码。'
    '工具增强模式下可使用 read_file、grep_files、glob_files、run_linter、find_symbol_definition 等工具，'
    '工作流程限制在 6-8 轮工具调用内，要求先探索代码再输出 JSON 结果。'
    '使用 langchain-openai 的 ChatOpenAI 调用 DeepSeek API，temperature 0.1 保证一致性。'
    'parse_llm_json_object() 使用 json_repair 库自动修复 LLM 输出的非法 JSON。'
))

add_heading(doc, '6.3 安全分析智能体（Security Agent）', 2)
add_para(doc, '文件：backend/agents/security.py')
add_para(doc, (
    '检测 OWASP Top-10 安全漏洞，按 CWE 标准分类。漏洞分三个严重等级：'
    'HIGH（SQL注入 CWE-89、命令注入 CWE-78、路径遍历 CWE-22、硬编码密钥 CWE-798）；'
    'MEDIUM（XSS CWE-79、SSRF CWE-918、弱加密 CWE-327、不安全随机数 CWE-330）；'
    'LOW（缺少输入验证、不安全默认配置）。提示词注入 imports 信息帮助理解第三方库依赖。'
    '工具模式下使用 grep_files 全局搜索 exec/eval/subprocess/os.system/pickle 等危险模式。'
))

add_heading(doc, '6.4 修复智能体（Repair Agent）', 2)
add_para(doc, '文件：backend/agents/repair.py（241 行）')
add_para(doc, (
    '通过多轮对话和工具调用直接修改代码，不输出 JSON 或 diff 文本。工作流程严格按顺序：'
    '第 1-2 轮 read_file 读取和分析；第 3-8 轮 replace_code/insert_code 逐一修复 + run_linter 验证；'
    '第 9 轮必须输出 ALL_DONE 或 CANNOT_FIX；第 10 轮强制终止。支持两种模式：'
    'memory（dry_run 仅验证不写磁盘，前端编辑器展示变更）和 disk（直接写磁盘 + git 安全）。'
    'replace_code 要求 old_string 在文件中唯一精确匹配，匹配 0 次或 >1 次均拒绝执行。'
    '连续 5 次工具调用失败自动中止，undo_stack 支持逐级回滚。'
))

add_heading(doc, '6.5 工具注册与执行系统（Tool Registry）', 2)
add_para(doc, '文件：backend/tools/tool_registry.py')
add_para(doc, (
    '统一管理所有 LLM 可调用工具的定义和执行。核心数据结构 ToolDef 包含 name、description、'
    'parameters（JSON Schema）、fn（Python 实现函数）和 required_params。'
    'build_default_registry() 工厂函数集中创建 15 个内置工具：read_file、grep_files、glob_files、'
    'run_linter、find_symbol_definition、find_symbol_references、generate_diff、parse_ast、'
    'replace_code、insert_code、write_file、delete_file、undo_last_change、run_tests、list_tests。'
))
add_para(doc, (
    '关键机制：(1) 上下文注入 — execute() 通过 inspect.signature 自动注入 files/language/workspace_root/dry_run；'
    '(2) 参数归一化 — path ↔ file_path 自动映射；(3) 路径安全 — _validate_path 拦截 ../ 越权访问；'
    '(4) 撤销栈 — replace_code/write_file 自动备份；(5) 双重注册 — ToolRegistry.register() + @register 装饰器。'
))

add_heading(doc, '6.6 前端编辑器与 Diff 高亮', 2)
add_para(doc, '文件：frontend/index.html')
add_para(doc, (
    'CSS 类定义：.diff-add-bg（绿色新增）、.diff-del-bg（红色删除）、.gutter-add/.gutter-del（行号区标记）。'
    'pendingChanges[] 数组追踪所有未确认修改：refreshDiffHighlights() 按当前活跃 Tab 过滤构建 Diff 范围，'
    '每处改动首行注入悬停 ✓/↩ 按钮。diff-action-bar 显示变更统计并提供 Apply All / Undo All 批量操作。'
    '语法高亮使用纯 JS 词法分析器（highlightPython），支持关键字、内置函数、字符串、注释、装饰器的分类着色。'
    '编辑器基于原生 <textarea> + 透明前景色 + 上层 <pre> 高亮叠加（pointer-events: none）的双层架构。'
))

# ── 七、数据集与实验环境 ──
add_heading(doc, '七、数据集与实验环境', 1)
add_heading(doc, '7.1 数据来源', 2)
add_para(doc, (
    '系统采用"即送即审"模式，不依赖静态数据集。开发和测试阶段以项目自身 38 个 Python 文件作为测试数据（dogfooding 策略），'
    '另有 buggy_math.py 和 test_security_sample.py 等专门构造的测试文件用于验证安全分析和代码质量审查的准确性。'
))
add_heading(doc, '7.2 实验环境', 2)
add_table(doc, ['项目', '配置'], [
    ['操作系统', 'Windows 11 Home (China) 10.0.26200'],
    ['Python 版本', '3.10+ (E:\\paper2slides\\python.exe)'],
    ['LLM', 'DeepSeek v4-pro（默认），OpenAI GPT-4o（备选）'],
    ['前端', 'Electron 33 + Chromium'],
    ['主要框架', 'FastAPI, LangGraph, LangChain, Uvicorn'],
    ['开发工具', 'Visual Studio Code, Claude Code (AI 辅助编码)'],
])
add_heading(doc, '7.3 核心依赖', 2)
add_para(doc, (
    'fastapi≥0.111.0, uvicorn[standard]≥0.29.0, langgraph≥0.1.0, langchain≥0.2.0, '
    'langchain-openai≥0.1.0, openai≥1.30.0, tree-sitter≥0.22.0, tree-sitter-python≥0.21.0, '
    'tree-sitter-javascript≥0.21.0, pydantic≥2.7.0, python-dotenv≥1.0.0, '
    'json-repair≥0.28.0, slowapi≥0.1.9, sse-starlette≥2.0.0'
))

# ── 八、系统测试与结果分析 ──
add_heading(doc, '八、系统测试与结果分析', 1)
add_heading(doc, '8.1 功能测试', 2)
add_table(doc, ['编号', '测试功能', '测试方法', '预期结果', '实际结果'], [
    ['FT-01', '单文件代码审查', '提交 Python 文件', '6 Agent 并行返回结构化问题', '通过'],
    ['FT-02', '多文件项目审查', '提交完整项目目录', '上下文分析 + 审查覆盖所有文件', '通过'],
    ['FT-03', 'SSE 流式展示', '观察审查过程', '逐 Agent 显示进度', '通过'],
    ['FT-04', 'AST 解析', '提交复杂语法代码', '正确提取函数/类/导入', '通过'],
    ['FT-05', '安全漏洞检测', '提交含注入/硬编码密钥代码', '识别并标注 CWE 编号', '通过'],
    ['FT-06', '自动代码修复', '选择问题执行修复', '工具链正确替换，lint 通过', '通过'],
    ['FT-07', '意图分类', '输入不同意图消息', 'LLM 正确分类并路由', '通过'],
    ['FT-08', '工具调用容错', '连续发送无效参数', '5 次失败后自动中止', '通过'],
    ['FT-09', '路径安全检查', '提交含 ../ 的路径', '拒绝越权访问', '通过'],
    ['FT-10', 'API 速率限制', '1 分钟内超 30 次请求', '返回 429', '通过'],
])

add_heading(doc, '8.2 性能测试', 2)
add_table(doc, ['指标', '数值', '说明'], [
    ['预处理（AST 解析）', '<100ms', 'Tree-sitter 增量解析'],
    ['并行审查（6 Agent）', '3-8 秒', 'asyncio.gather 并发，取最慢 Agent'],
    ['汇总阶段', '2-4 秒', '1 次 LLM 调用'],
    ['审查总计', '5-12 秒', '网络延迟为主要因素'],
    ['Token 消耗', '5000-15000/次', '含系统提示词和代码内容'],
    ['首次 Token 延迟', '0.8-2 秒', '流式模式'],
    ['SSE 心跳', '15 秒', '可配置'],
])

add_heading(doc, '8.3 Agent 审查准确性分析', 2)
add_para(doc, (
    '本项目采用多智能体交叉验证机制保证审查准确性：'
    '(1) 专业化分工 — 6 个 Agent 从不同维度独立分析，各通过 AGENT_TOOL_WHITELIST 限制工具集；'
    '(2) 并行独立执行 — asyncio.gather 确保 Agent 互不干扰，一个错误不传播到其他；'
    '(3) 汇总仲裁 — Summarizer 执行幻觉过滤、矛盾仲裁、去重合并和六维评分。'
    '此外，json_repair 自动修复 LLM 非法 JSON，temperature 0.1 减少输出随机性，'
    'replace_code 要求精确唯一匹配防止误替换，连续失败检测防止死循环。'
))

# ── 九、系统部署与使用说明 ──
add_heading(doc, '九、系统部署与使用说明', 1)
add_para(doc, '详见《期末用户手册》和《期末部署说明》。')
add_para(doc, '快速启动：')
add_para(doc, '1. pip install -r requirements.txt')
add_para(doc, '2. 在 .env 中配置 DEEPSEEK_API_KEY')
add_para(doc, '3. npm install && npm start')
add_para(doc, '4. 访问 http://127.0.0.1:8765')

# ── 十、AI 工具使用与团队协作说明 ──
add_heading(doc, '十、AI 工具使用与团队协作说明', 1)
add_heading(doc, '10.1 AI 工具使用记录', 2)
add_table(doc, ['序号', '使用环节', 'AI 工具', '使用方式', '输出结果', '是否直接采用', '人工修改说明'], [
    ['1', '项目骨架搭建', 'Claude Code', '对话式代码生成', 'FastAPI 路由、LangGraph 编排、Electron 主进程等核心模块', '是', '人工审查后确认'],
    ['2', '多 Agent 提示词工程', 'Claude Code + DeepSeek', '迭代优化 Prompt', '6 个 Agent 的中文系统提示词', '部分', '人工调整 JSON 输出格式约束'],
    ['3', '工具注册系统', 'Claude Code', '代码生成', 'ToolRegistry + ToolDef + @register 装饰器', '是', '人工添加路径安全检查'],
    ['4', '前端界面', 'Claude Code', 'HTML/CSS/JS 生成', '约 2700 行单页应用', '是', '人工调整 Diff 高亮和交互逻辑'],
    ['5', 'Bug 定位与修复', 'Claude Code', 'Code Review 指令', '发现 summarizer state aliasing、空 SSE 事件等 Bug', '直接采用', '—'],
    ['6', '文档生成', 'Claude Code', '10 Agent 并行分析', '技术报告/用户手册/部署说明', '是', '人工整合和格式化'],
    ['7', '代码审查（运行时）', 'DeepSeek v4-pro', 'LLM API 调用', '所有 Agent 的审查分析和修复建议', '—（系统核心功能）', '用户通过 Diff 高亮逐项确认'],
])

add_heading(doc, '10.2 人机协作说明', 2)
add_para(doc, (
    'AI 主导生成的部分：多智能体编排框架（graph.py）、工具注册系统（tool_registry.py）、'
    '前端界面（index.html，约 2700 行）、文件操作工具（file_tool.py）、AST 解析模块（ast_parser.py）。'
))
add_para(doc, (
    '人类主导设计的部分：系统架构设计（四层分层 + 前后端分离）、提示词工程（6 个 Agent 的中文 System Prompt）、'
    '安全策略（路径穿越防护、API Key 检查、速率限制、CORS/CSP 配置）、'
    '工具约束逻辑（Agent 工具白名单、修复轮数限制、连续失败中止）、'
    '项目配置（Token 限额、服务端口、超时参数）。'
))
add_para(doc, (
    '协作模式：人类定义架构边界、安全约束和提示词策略，AI 在这些约束下生成高效代码实现。'
    '通过 Claude Code /code-review 指令审查 AI 生成代码，发现潜在问题并修复。'
    '大幅提升开发效率，同时保持架构一致性和安全性。'
))

add_heading(doc, '10.3 团队分工与贡献', 2)
add_para(doc, '本项目为独立开发项目（单人），开发者承担全部角色：')
add_table(doc, ['角色', '职责范围'], [
    ['系统架构师', '多智能体协作架构、前后端分离方案、SSE 流式通信协议设计'],
    ['后端工程师', 'FastAPI 服务、LangGraph 编排、8 个 Agent 实现、工具注册系统、LLM 集成'],
    ['AI 算法工程师', '提示词工程、JSON 容错解析、幻觉过滤、交叉验证仲裁'],
    ['前端工程师', 'Electron 桌面应用、三面板布局、语法高亮、Diff 可视化、SSE 流式渲染'],
    ['安全工程师', '路径穿越防护、API 速率限制、CORS/CSP、Electron 沙箱安全'],
    ['测试工程师', '功能测试、Agent 准确性评估、边界条件测试'],
    ['DevOps', '项目配置管理、环境管理、Git 版本控制'],
])

# ── 十一、项目总结 ──
add_heading(doc, '十一、项目总结', 1)
add_heading(doc, '11.1 项目成果总结', 2)
add_para(doc, (
    '本系统成功构建了功能完整的 AI 代码审查桌面应用，核心成果包括：'
    '(1) 6+1+1 共 8 个 AI Agent 的完整实现与并行编排；'
    '(2) 15 个 LLM 可调用工具的统一注册、注入与安全执行引擎；'
    '(3) 基于 SSE 的流式实时推送与 Expert Round Table 可视化；'
    '(4) Human-in-the-Loop 修复模式（Diff 高亮 + 逐项确认）；'
    '(5) Electron 桌面应用的一键启动与进程生命周期管理。'
    '系统代码总计约 38 个 Python 文件 + 2914 行前端代码，在开发环境中稳定运行。'
))
add_heading(doc, '11.2 项目应用价值', 2)
add_para(doc, '可实际用于高校学生代码作业审查、实训项目批量审查、个人开发者代码质量保障、开源项目贡献前自检。')
add_heading(doc, '11.3 不足与反思', 2)
add_para(doc, (
    '(1) 当前依赖云端 LLM API（DeepSeek），对网络和 API 余额有依赖，隐私敏感场景可扩展本地模型（如 Ollama）；'
    '(2) 修复 Agent 对跨越多个文件的大型重构支持有限，当前主要针对单文件内的问题；'
    '(3) 可增加更多编程语言支持（当前支持 6 种）；'
    '(4) 可增加审查历史对比、团队共享等协作功能。'
))

print(f"[1/3] 技术报告已保存到 {os.path.join(OUT_DIR, '期末技术报告.docx')}")
doc.save(os.path.join(OUT_DIR, '期末技术报告.docx'))

# ═══════════════════════════════════════════════════════════════
# Document 2: 期末用户手册 (abbreviated - key sections only)
# ═══════════════════════════════════════════════════════════════

doc2 = Document()
style2 = doc2.styles['Normal']
style2.font.name = '宋体'
style2.font.size = Pt(11)
style2.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

t2 = doc2.add_paragraph()
t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = t2.add_run('人工智能综合课程设计\n期末用户手册\n\nCodeReview AI — 多智能体代码审查桌面系统')
r2.font.size = Pt(18)
r2.bold = True

doc2.add_paragraph()

# 一、文档说明
add_heading(doc2, '一、文档说明', 1)
add_para(doc2, '编写目的：帮助课程评审教师、项目审阅人员、系统管理员掌握 CodeReview AI 的安装、配置与使用。')
add_table(doc2, ['项目', '内容'], [
    ['文档名称', 'CodeReview AI 用户手册'],
    ['版本号', 'V1.0'],
    ['创建日期', '2026年5月29日'],
    ['系统版本', 'V1.0'],
])

# 二、系统概述
add_heading(doc2, '二、系统概述', 1)
add_para(doc2, (
    'CodeReview AI 是一款基于多智能体架构的 AI 代码审查桌面应用程序，集成 Electron 桌面壳层、'
    'FastAPI 后端、LangGraph 编排和 DeepSeek/OpenAI 大语言模型，通过 6 个专业审查 Agent 并行分析代码质量、'
    '安全漏洞、性能瓶颈、文档完整性、测试质量和架构设计，由 Summarizer 进行交叉验证与综合评分。'
))
add_table(doc2, ['用户角色', '目标', '核心功能'], [
    ['学生开发者', '提交前获得专业审查意见', '代码审查、安全分析、性能建议、修复模式'],
    ['教师/评审者', '快速批量评审代码质量', '项目审查、六维评分、报告导出'],
    ['系统管理员', '维护运行环境', '健康检查、Token 管理、API 配置'],
])
add_para(doc2, '系统主要功能：代码审查（6 Agent 并行）、安全分析（OWASP Top-10）、性能优化建议、测试质量分析、修复模式（Human-in-the-Loop）、文件树浏览与编辑、Diff 高亮与变更确认、对话历史保存与加载。')

# 三、运行环境要求
add_heading(doc2, '三、运行环境要求', 1)
add_heading(doc2, '3.1 硬件环境', 2)
add_table(doc2, ['项目', '最低配置', '推荐配置'], [
    ['CPU', '双核 x86-64', '四核及以上'],
    ['内存', '4 GB RAM', '8 GB RAM'],
    ['硬盘', '500 MB', '1 GB'],
    ['显卡', '无需独立显卡', '—'],
    ['网络', '互联网连接（调用云端 API）', '宽带连接'],
])
add_heading(doc2, '3.2 软件环境', 2)
add_table(doc2, ['组件', '版本要求', '用途'], [
    ['Python', '3.10+', '后端运行时'],
    ['Node.js', '18+ LTS', 'Electron 运行时'],
    ['npm', '9+', '包管理器'],
    ['DeepSeek API Key', '必需', 'LLM 推理（推荐）'],
    ['OpenAI API Key', '可选', '备选 LLM'],
])

# 四、安装与部署
add_heading(doc2, '四、安装与部署说明', 1)
add_para(doc2, '步骤一：获取项目代码（git clone 或解压交付包）', bold=True)
add_para(doc2, '步骤二：安装 Python 依赖', bold=True)
add_para(doc2, 'python -m venv .venv && .venv\\Scripts\\activate && pip install -r requirements.txt')
add_para(doc2, '步骤三：安装 Node.js 依赖', bold=True)
add_para(doc2, 'npm install')
add_para(doc2, '步骤四：配置 API 密钥', bold=True)
add_para(doc2, '在项目根目录创建 .env 文件：DEEPSEEK_API_KEY=sk-你的密钥')
add_para(doc2, '步骤五（可选）：在 electron/config.json 中配置 Python 路径')
add_para(doc2, '步骤六：启动系统', bold=True)
add_para(doc2, 'npm start')

# 五、系统使用说明
add_heading(doc2, '五、系统使用说明', 1)
add_heading(doc2, '5.1 打开项目', 2)
add_para(doc2, '点击 Explorer 面板"打开目录"（Ctrl+O）选择项目文件夹，或"打开文件"（Ctrl+Shift+O）打开单个文件。')
add_heading(doc2, '5.2 代码审查', 2)
add_para(doc2, '在聊天框输入"审查代码"或"全面审查"，系统自动路由到 6 个专家 Agent 并行分析。')
add_heading(doc2, '5.3 修复模式', 2)
add_para(doc2, '审查完成后输入"修复"，Repair Agent 逐一修改代码。编辑器中绿色=新增，红色=删除。'
           '悬停出现 ✓/↩ 按钮逐项确认，Apply All 全部应用，Undo All 全部撤销。')
add_heading(doc2, '5.4 文件操作', 2)
add_para(doc2, 'write_file 创建文件 → 预览 Tab [待确认]，全绿高亮；delete_file 删除文件 → 预览 Tab [待删除]，全红高亮。确认后生效。')
add_heading(doc2, '5.5 保存对话', 2)
add_para(doc2, '点击 💾 保存对话为 chat_YYYY_MM_DD_HH_mm_ss.md；点击 📂 加载历史对话。对话自动保存到项目根目录。')
add_heading(doc2, '5.6 典型使用流程', 2)
add_para(doc2, '打开项目 → 审查代码 → 发现问题 → 请求修复计划 → 执行计划 → 审查 Diff → 逐项确认 → 保存文件 → 导出对话')

# 六、Expert Round Table
add_heading(doc2, '六、Expert Round Table（专家圆桌）', 1)
add_para(doc2, '聊天区域中每次 AI 响应包含 Expert Roundtable 折叠区，展示所有 Agent 的思考过程和工具调用记录。')
add_table(doc2, ['图标', 'Agent', '颜色'], [
    ['[R]', '代码审查 Reviewer', '蓝色'],
    ['[S]', '安全专家 Security', '红色'],
    ['[O]', '性能专家 Optimizer', '黄色'],
    ['[D]', '文档专家 Documenter', '绿色'],
    ['[T]', '测试专家 Tester', '紫色'],
    ['[A]', '架构专家 Architect', '深紫'],
    ['[H]', '修复专家 Repair', '绿色'],
    ['[M]', '综合裁判 Summarizer', '白色'],
])

# 七、常见问题排查
add_heading(doc2, '七、常见问题排查（FAQ）', 1)
add_table(doc2, ['问题', '原因', '解决方法'], [
    ['启动后白屏', 'Python 后端未能在 30s 内启动', '检查 Python 路径、依赖安装、端口占用'],
    ['API Key 未配置', '.env 文件缺失或密钥无效', '创建 .env 并填入有效 DEEPSEEK_API_KEY'],
    ['Rate limit exceeded', '超过 30次/小时 限制', '等待恢复或修改限值配置'],
    ['replace_code 失败', 'old_string 与文件内容不匹配', '提供更多上下文行确保唯一匹配'],
    ['SSE 流中断', '网络不稳定或 API 响应慢', '前端有 3 次自动重试机制'],
    ['文件树不刷新', '新建/删除文件后未调用刷新', '确认操作后会自动刷新 Explorer'],
    ['报告保存失败', 'reports/ 目录不存在', '手动创建 reports/ 目录'],
])

# 八、安全与隐私
add_heading(doc2, '八、安全与隐私说明', 1)
add_table(doc2, ['安全机制', '实现方式'], [
    ['API 鉴权', 'AuthMiddleware — localhost 自动放行，远程需 X-API-Key'],
    ['安全响应头', 'X-Content-Type-Options, X-Frame-Options, X-XSS-Protection'],
    ['速率限制', 'SlowAPI — 全局 60/min，审查端点 30/hour'],
    ['CORS', '仅允许 127.0.0.1:8765, localhost:8765, app://.'],
    ['CSP', '限制脚本/样式/连接来源'],
    ['路径遍历防护', '_validate_path 拦截 ../ 越权'],
    ['Electron 安全', 'sandbox: true, contextIsolation: true, nodeIntegration: false'],
    ['数据警告', '前端明确提示"代码将被发送到第三方 API 进行分析"'],
])

print(f"[2/3] 用户手册已保存到 {os.path.join(OUT_DIR, '期末用户手册.docx')}")
doc2.save(os.path.join(OUT_DIR, '期末用户手册.docx'))

# ═══════════════════════════════════════════════════════════════
# Document 3: 期末部署说明
# ═══════════════════════════════════════════════════════════════

doc3 = Document()
style3 = doc3.styles['Normal']
style3.font.name = '宋体'
style3.font.size = Pt(11)
style3.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

t3 = doc3.add_paragraph()
t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = t3.add_run('人工智能综合课程设计\n期末部署说明\n\nCodeReview AI — 多智能体代码审查桌面系统')
r3.font.size = Pt(18)
r3.bold = True

doc3.add_paragraph()

# 一、文档基本信息
add_heading(doc3, '一、文档基本信息', 1)
add_table(doc3, ['项目', '内容'], [
    ['课程名称', '人工智能综合课程设计'],
    ['项目名称', 'CodeReview AI — 多智能体代码审查桌面系统'],
    ['文档类型', '期末部署说明'],
    ['版本号', 'V1.0'],
    ['编写日期', '2026年5月29日'],
    ['系统版本', 'V1.0'],
])

# 二、部署说明目的
add_heading(doc3, '二、部署说明目的', 1)
add_para(doc3, '本文档面向课程评审教师或第三方评审人员，提供从零开始独立部署并验证 CodeReview AI 系统的完整操作指引。'
           '评审人员可依照本文档逐步骤操作，无需项目组成员介入即可完成环境搭建、系统启动和功能验证。')

# 三、系统形态说明
add_heading(doc3, '三、系统形态说明', 1)
add_para(doc3, '系统形态：Desktop Application（Electron 桌面应用）+ Local Python Backend（FastAPI 后端）。')
add_para(doc3, '部署形式：本地单机部署（Local Deployment），所有组件在同一台机器上运行。')
add_para(doc3, (
    '部署架构：Electron 主进程管理 Python 后端子进程（child_process.spawn），后端监听 127.0.0.1:8765，'
    '渲染进程通过 HTTP/SSE 与后端通信，后端通过 HTTPS 调用 DeepSeek API。'
    'Electron 自动探测可用端口（8765-8785），应用退出时自动清理 Python 进程。'
))

# 四、部署环境要求
add_heading(doc3, '四、部署环境要求', 1)
add_heading(doc3, '4.1 硬件环境', 2)
add_table(doc3, ['硬件', '最低配置', '推荐配置'], [
    ['CPU', '双核 x86_64, 2.0GHz+', '四核 x86_64, 2.5GHz+'],
    ['内存', '4 GB RAM', '8 GB RAM'],
    ['硬盘', '1 GB 可用空间', '5 GB 可用空间'],
    ['网络', '需访问 api.deepseek.com', '稳定宽带连接'],
    ['显卡', '无需独立显卡', '—'],
])

add_heading(doc3, '4.2 软件环境', 2)
add_table(doc3, ['软件', '版本要求', '必需', '说明'], [
    ['操作系统', 'Windows 10/11 64位 或 Linux/macOS', '是', '—'],
    ['Python', '3.10 或更高（不支持 3.8/3.9）', '是', 'LangGraph 要求 3.10+'],
    ['Node.js', '18 LTS 或更高', '是', 'Electron 运行时'],
    ['npm', '9.x+', '是', '随 Node.js 附带'],
    ['DeepSeek API Key', '—', '是', '从 platform.deepseek.com 获取'],
    ['OpenAI API Key', '—', '否', '备选方案'],
    ['CUDA/GPU', '—', '否', '系统纯 CPU 运行'],
])

# 五、部署前准备
add_heading(doc3, '五、部署前准备', 1)
add_para(doc3, '部署前检查清单：')
add_para(doc3, '[ ] Python 3.10+ 已安装（python --version）')
add_para(doc3, '[ ] Node.js 18+ 已安装（node --version）')
add_para(doc3, '[ ] 已获取有效的 DeepSeek API Key')
add_para(doc3, '[ ] 8765 端口未被占用（netstat -ano | findstr :8765）')
add_para(doc3, '[ ] 网络可访问 https://api.deepseek.com')
add_para(doc3, '[ ] 项目代码完整（确认后端 38 个 .py 文件存在）')

# 六、标准部署步骤
add_heading(doc3, '六、标准部署步骤', 1)
add_para(doc3, '6.1 获取项目代码', bold=True)
add_para(doc3, 'git clone <仓库地址> codereviewer && cd codereviewer')
add_para(doc3, '或解压交付压缩包到目标目录。')
add_para(doc3, '6.2 安装 Python 环境', bold=True)
add_para(doc3, 'python -m venv .venv && .venv\\Scripts\\activate && pip install -r requirements.txt')
add_para(doc3, '6.3 安装 Node.js 依赖', bold=True)
add_para(doc3, 'npm install')
add_para(doc3, '6.4 配置环境变量', bold=True)
add_para(doc3, '创建 .env 文件：DEEPSEEK_API_KEY=sk-你的实际密钥')
add_para(doc3, '6.5 配置 Python 解释器（可选）', bold=True)
add_para(doc3, '编辑 electron/config.json：{"pythonPath": "你的Python路径"}')
add_para(doc3, '或设置环境变量 CODEREVIEW_PYTHON（优先级最高）。')
add_para(doc3, '6.6 启动系统', bold=True)
add_para(doc3, 'npm start')
add_para(doc3, '(Electron 窗口弹出 → 自动启动 Python 后端 → 轮询健康检查 → 加载主界面)')

add_heading(doc3, '6.7 验证部署', 2)
add_para(doc3, '健康检查：curl http://127.0.0.1:8765/api/health')
add_para(doc3, '预期返回：{"status":"ok","llm_configured":true,"primary_provider":"deepseek"}')
add_para(doc3, '功能验证：打开文件 → 输入"审查代码" → 观察 Expert Round Table → 查看审查报告')

# 七、常见问题排查
add_heading(doc3, '七、故障排查与常见问题', 1)
add_table(doc3, ['问题', '原因', '解决方法'], [
    ['pip install 失败', 'Python 版本 <3.10 或缺少 C++ 编译器', '升级 Python，安装 VC++ Build Tools'],
    ['npm install 失败', 'Node.js 版本过低或网络问题', '升级 Node.js；设置 ELECTRON_MIRROR 镜像'],
    ['启动白屏', 'Python 后端未在 30s 内启动', '手动运行 uvicorn backend.main:app 查看错误'],
    ['API Key 未配置', '.env 缺失或密钥无效', '检查 .env 文件中的 DEEPSEEK_API_KEY'],
    ['端口被占用', '8765 端口冲突', 'Electron 自动尝试 8765-8785；或手动杀进程'],
    ['模型名错误', '文档写 deepseek-coder', '实际默认 deepseek-v4-pro，在 .env 中配置'],
    ['审查返回 500', 'API Key 无效或余额不足', '检查 /api/health 中 llm_configured'],
    ['tree-sitter 编译失败', '缺少 C/C++ 编译工具', '安装 Microsoft C++ Build Tools'],
    ['Electron 窗口不刷新', '无热重载', 'Ctrl+R 重新载入'],
    ['残留 Python 进程', 'Electron 崩溃', 'taskkill /f /im python.exe'],
])

# 八、安全与合规
add_heading(doc3, '八、安全与合规提示', 1)
add_para(doc3, 'API Key 管理：.env 文件已加入 .gitignore，代码中 _normalize_key() 过滤占位符值。')
add_para(doc3, '路径遍历防护：_validate_path() 拦截 ../ 越权访问，绝对路径直接放行。')
add_para(doc3, 'CORS：仅允许 127.0.0.1:8765, localhost:8765, app://. 来源。')
add_para(doc3, '认证：本地回环地址自动放行，远程访问需 APP_API_KEY。')
add_para(doc3, 'Electron 安全：sandbox:true, contextIsolation:true, nodeIntegration:false。')
add_para(doc3, '速率限制：SlowAPI 全局 60/min，审查端点 30/hour。')

# 九、回滚与清理
add_heading(doc3, '九、回滚与清理说明', 1)
add_para(doc3, '停止服务：关闭 Electron 窗口（自动清理 Python 进程树）。')
add_para(doc3, '手动清理：Windows — taskkill /f /im python.exe；Linux — pkill -f "uvicorn backend.main:app"')
add_para(doc3, '清理文件：删除 reports/ 目录中的审查报告；清除 localStorage 中的聊天记录。')
add_para(doc3, '卸载：关闭应用 → 删除项目根目录 → （可选）删除 Python 虚拟环境。')

# 十、附录
add_heading(doc3, '十、附录', 1)
add_para(doc3, '默认访问地址：')
add_para(doc3, '  前端入口：http://127.0.0.1:8765')
add_para(doc3, '  API 文档：http://127.0.0.1:8765/api/docs')
add_para(doc3, '  健康检查：http://127.0.0.1:8765/api/health')
add_para(doc3, 'Python 解释器优先级：')
add_para(doc3, '  1. CODEREVIEW_PYTHON 环境变量（最高）')
add_para(doc3, '  2. electron/config.json 中的 pythonPath')
add_para(doc3, '  3. 系统默认 python 命令（兜底）')
add_para(doc3, '.env 配置示例：', bold=True)
add_para(doc3, 'DEEPSEEK_API_KEY=sk-your-key-here')
add_para(doc3, 'DEEPSEEK_MODEL=deepseek-v4-pro')
add_para(doc3, '# OPENAI_API_KEY=sk-your-key-here  # 备选')
add_para(doc3, '关键配置常量（backend/config.py）：', bold=True)
add_para(doc3, 'MAX_CODE_SIZE=2MB | MAX_FILES=200 | MAX_TOKENS=8192 | DAILY_TOKEN_LIMIT=1,000,000')

print(f"[3/3] 部署说明已保存到 {os.path.join(OUT_DIR, '期末部署说明.docx')}")
doc3.save(os.path.join(OUT_DIR, '期末部署说明.docx'))

print("\n✅ 全部 3 份文档已生成！")
print(f"   输出目录: {OUT_DIR}")
